import requests
import json
from datetime import datetime, timedelta
from pathlib import Path
import re
from docx import Document


class FeishuDocumentAPI:
    def __init__(
        self,
        notice_dir,
        notice_message_heading,
        document_id,
        app_id,
        app_secret,
        message_heading_text,
        debug_offline_all_document_comments_response_json=False,
        save_all_document_blocks_response_as_json=True,
        debug_offline_all_document_blocks_response_json=False,
        save_all_document_comments_response_as_json=True,
    ) -> None:
        self.notice_dir = Path(notice_dir)
        self.notice_path = Path()
        script_dir = Path(__file__).resolve().parent
        self.notice_message_heading = notice_message_heading
        self.document_id = document_id
        tenant_access_token = self.obtain_tenant_access_token(app_id, app_secret)
        self.access_token = tenant_access_token
        self.str_specific_time_yesterday = "16-00"
        # Process all document comments firstly.
        # Unsolved comments need to be updated to the document.
        self.all_document_comments_response = None
        self.unsolved_document_comments_list = []
        self.preprocess_all_document_comments(
            debug_offline_all_document_comments_response_json
        )
        all_document_comments_response_path = (
            script_dir / "all_document_comments_response.json"
        )
        if save_all_document_comments_response_as_json is True:
            with open(all_document_comments_response_path, "w+", encoding="utf8") as f:
                json.dump(self.all_document_comments_response, f, ensure_ascii=False)
        # Process all document blocks.
        self.all_document_blocks_response = None
        self.all_document_blocks = []
        # The first block is the document title.
        self.all_document_children_block_ids = []
        self.message_heading_text = message_heading_text
        self.message_heading_block_id = None
        self.item_message_heading_block_index = None
        self.children_message_heading_block_index = None
        self.item_message_block_start_index = None
        self.children_message_block_start_index = None
        self.message_blocks_list = []
        self.preprocess_all_document_blocks(
            debug_offline_all_document_blocks_response_json
        )
        all_document_blocks_response_path = (
            script_dir / "all_document_blocks_response.json"
        )
        if save_all_document_blocks_response_as_json is True:
            with open(all_document_blocks_response_path, "w+", encoding="utf8") as f:
                json.dump(self.all_document_blocks_response, f, ensure_ascii=False)

    def obtain_tenant_access_token(self, app_id, app_secret):
        url = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal"
        headers = {
            "Content-Type": "application/json; charset=utf-8",
        }
        payload = json.dumps(
            {
                "app_id": app_id,
                "app_secret": app_secret,
            }
        )
        response = requests.request("POST", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        return response_json["tenant_access_token"]

    """
    Methods for comments.
    """

    def obtain_all_document_comments(self):
        url = (
            "https://open.feishu.cn/open-apis/drive/v1/files/"
            + self.document_id
            + "/comments?file_type=docx"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("GET", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def delete_a_reply(self, document_comment):
        # Deprecated
        # The document app can only delete replies created by itself.
        reply_id = document_comment["reply_list"]["replies"][0]["reply_id"]
        url = (
            "https://open.feishu.cn/open-apis/drive/v1/files/"
            + self.document_id
            + "/comments/"
            + document_comment["comment_id"]
            + "/replies/"
            + reply_id
            + "?file_type=docx"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("DELETE", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def solve_a_reply(self, document_comment):
        url = (
            "https://open.feishu.cn/open-apis/drive/v1/files/"
            + self.document_id
            + "/comments/"
            + document_comment["comment_id"]
            + "?file_type=docx"
        )
        payload = json.dumps({"is_solved": True})  # type: ignore
        access_token = "Bearer " + self.access_token
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": access_token,
        }
        response = requests.request("PATCH", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def delete_a_reply(self, document_comment):
        url = (
            "https://open.feishu.cn/open-apis/drive/v1/files/"
            + self.document_id
            + "/comments/"
            + document_comment["comment_id"]
            + "/replies/"
            + document_comment["comment_id"]
            + "?file_type=docx"
        )
        payload = json.dumps({})
        access_token = "Bearer " + self.access_token
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": access_token,
        }
        response = requests.request("DELETE", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)

        reply = document_comment["reply_list"]["replies"][0]
        document_comment_text = reply["content"]["elements"][0]["text_run"][
            "text"
        ].replace("\n", " ")
        if response_json["code"] == 0:
            print(f"Sucessfully delete the reply: {document_comment_text}")

        return response_json

    def preprocess_all_document_comments_archived(
        self, debug_offline_all_document_comments_json
    ):
        if debug_offline_all_document_comments_json is True:
            script_dir = Path(__file__).resolve().parent
            all_document_comments_response_json_path = (
                script_dir / "all_document_comments_response.json"
            )
            with open(
                all_document_comments_response_json_path, "r", encoding="utf8"
            ) as f:
                self.all_document_comments_response = json.load(f)
        else:
            self.all_document_comments_response = self.obtain_all_document_comments()
        self.all_document_comments = self.all_document_comments_response["data"][
            "items"
        ]
        for document_comment in self.all_document_comments:
            # Process unsolved comments only.
            if document_comment["solver_user_id"] is None:
                self.unsolved_document_comments_list.append(document_comment)
            else:
                continue
        print(
            "{} unsolved document comment(s) obtained".format(
                len(self.unsolved_document_comments_list)
            )
        )
        for unsolved_document_comment in self.unsolved_document_comments_list:
            # Process unsolved comments only.
            reply = unsolved_document_comment["reply_list"]["replies"][0]
            unsolved_document_comment_text = reply["content"]["elements"][0][
                "text_run"
            ]["text"].replace("\n", " ")
            try:
                create_blocks_response_json = self.create_blocks(
                    unsolved_document_comment_text
                )
                if create_blocks_response_json["code"] == 0:
                    create_blocks_response_content = create_blocks_response_json[
                        "data"
                    ]["children"][0]["text"]["elements"][0]["text_run"]["content"]
                    print(
                        "sucessfully created a document children block based on the comment:",
                        create_blocks_response_content,
                    )
                    try:
                        solve_a_reply_response_json = self.solve_a_reply(
                            document_comment
                        )
                        if solve_a_reply_response_json["code"] == 0:
                            print(
                                "sucessfully solve the comment:",
                                unsolved_document_comment_text,
                            )
                    except Exception as e:
                        print(f"A error occurred when calling delete_a_reply: {e}")
            except Exception as e:
                print(f"An error occurred when calling create_blocks: {e}")

    def preprocess_all_document_comments(
        self, debug_offline_all_document_comments_json
    ):
        def is_comment_solved_48_hours_before(document_comment):
            if document_comment["update_time"] is None:
                return False
            else:
                document_comment_update_time_timestamp_s = document_comment[
                    "update_time"
                ]
                document_comment_update_time = datetime.fromtimestamp(
                    document_comment_update_time_timestamp_s
                )
                current_time = datetime.now()
                time_difference = current_time - document_comment_update_time
                time_difference_seconds = time_difference.total_seconds()
                if 48 * 60 * 60 < time_difference_seconds:
                    return True
                else:
                    return False

        if debug_offline_all_document_comments_json is True:
            script_dir = Path(__file__).resolve().parent
            all_document_comments_response_json_path = (
                script_dir / "all_document_comments_response.json"
            )
            with open(
                all_document_comments_response_json_path, "r", encoding="utf8"
            ) as f:
                self.all_document_comments_response = json.load(f)
        else:
            self.all_document_comments_response = self.obtain_all_document_comments()
            self.all_document_comments = self.all_document_comments_response["data"][
                "items"
            ]

        for document_comment in self.all_document_comments:
            # Delete comments solved 48 hours ago
            if is_comment_solved_48_hours_before(document_comment):
                self.delete_a_reply(document_comment)
            # Process unsolved comments only.
            if document_comment["solver_user_id"] is None:
                self.unsolved_document_comments_list.append(document_comment)
            else:
                continue
        print(
            "{} unsolved document comment(s) obtained".format(
                len(self.unsolved_document_comments_list)
            )
        )

        if len(self.unsolved_document_comments_list) != 0:
            self.check_notice_exists()
            notice = Document(self.notice_path)
            target_paragraph_text = self.notice_message_heading

            target_paragraph = None
            for paragraph_index, paragraph in enumerate(notice.paragraphs):
                if target_paragraph_text in paragraph.text:
                    target_paragraph = paragraph
                    break

            for unsolved_document_comment in self.unsolved_document_comments_list:
                # Process unsolved comments only.
                reply = unsolved_document_comment["reply_list"]["replies"][0]
                unsolved_document_comment_text = reply["content"]["elements"][0][
                    "text_run"
                ]["text"].replace("\n", " ")
                # Deliver messages
                new_paragraph = notice.add_paragraph(unsolved_document_comment_text)
                # notice.paragraphs.insert(paragraph_index + 1, new_paragraph)
                target_paragraph._p.addnext(new_paragraph._p)
                notice.save(self.notice_path)
                try:
                    now = datetime.now()
                    formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
                    # print(f"formatted_time: {formatted_time}")
                    block_creation_text_content = (
                        formatted_time + "【已通知】" + unsolved_document_comment_text
                    )
                    create_blocks_response_json = self.create_blocks(
                        block_creation_text_content
                    )
                    if create_blocks_response_json["code"] == 0:
                        create_blocks_response_content = create_blocks_response_json[
                            "data"
                        ]["children"][0]["text"]["elements"][0]["text_run"]["content"]
                        print(
                            "Sucessfully created a document children block based on the comment:",
                            create_blocks_response_content,
                        )
                        try:
                            solve_a_reply_response_json = self.solve_a_reply(
                                document_comment
                            )
                            if solve_a_reply_response_json["code"] == 0:
                                print(
                                    "sucessfully solve the comment:",
                                    unsolved_document_comment_text,
                                )
                        except Exception as e:
                            print(f"A error occurred when calling delete_a_reply: {e}")
                    else:
                        print(
                            f"A error occurred when creating a block based on the comment: {block_creation_text_content}"
                        )
                except Exception as e:
                    print(f"An error occurred when calling create_blocks: {e}")

    """'
    Methods for document blocks.
    """

    def obtain_plain_document_text_content(self):
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/raw_content"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("GET", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def obtain_all_document_blocks(self):
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks"
        )
        payload = ""
        access_token = "Bearer " + self.access_token
        headers = {"Authorization": access_token}
        response = requests.request("GET", url, headers=headers, data=payload)
        # print(response.text)
        response_json = json.loads(response.text)
        return response_json

    def preprocess_all_document_blocks(
        self, debug_offline_all_document_blocks_response_json
    ):
        if debug_offline_all_document_blocks_response_json is True:
            script_dir = Path(__file__).resolve().parent
            all_document_blocks_response_json_path = (
                script_dir / "all_document_blocks_response.json"
            )
            with open(
                all_document_blocks_response_json_path, "r", encoding="utf8"
            ) as f:
                self.all_document_blocks_response = json.load(f)
        else:
            self.all_document_blocks_response = self.obtain_all_document_blocks()

        self.all_document_blocks = self.all_document_blocks_response["data"]["items"]
        self.all_document_children_block_ids = self.all_document_blocks[0]["children"]
        for block_index, block in enumerate(self.all_document_blocks):
            if (
                block["block_type"] == 3
                and block["heading1"]["elements"][0]["text_run"]["content"]
                == self.message_heading_text
            ):
                self.message_heading_block_id = block["block_id"]
                # print(
                #     f"{self.message_heading_text} block_id: ",
                #     self.message_heading_block_id,
                # )
                self.item_message_heading_block_index = block_index
                # print(
                #     f"{self.message_heading_text} item_block_index: ",
                #     self.item_message_heading_block_index,
                # )
                # The message_blocks_list does not contain the message title.
                self.message_blocks_list = self.all_document_blocks[
                    self.item_message_heading_block_index + 1 :
                ]
                break
        # The first message block index is 1 block(s) under the `message_heading_text``.
        self.item_message_block_start_index = self.item_message_heading_block_index + 1

        # The first block is the document title (parent block) which could not be deleted.
        # Only the children blocks could be deleted.
        # The index is calculated based on children blocks.
        for document_children_block_id_index, document_children_block_id in enumerate(
            self.all_document_children_block_ids
        ):
            if document_children_block_id == self.message_heading_block_id:
                self.children_message_heading_block_index = (
                    document_children_block_id_index
                )
                # print(
                #     f"{self.message_heading_text} children_block_index: ",
                #     self.children_message_heading_block_index,
                # )
                self.children_message_block_start_index = (
                    self.children_message_heading_block_index + 1
                )
                break

    def create_blocks(self, text_content):
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks/"
            + self.document_id
            + "/children?document_revision_id=-1"
        )
        payload = json.dumps(
            {
                "children": [
                    {
                        "block_type": 2,
                        "text": {
                            "elements": [
                                {
                                    "text_run": {
                                        "content": text_content,
                                        "text_element_style": {},
                                    }
                                },
                            ],
                            "style": {},
                        },
                    }
                ],
                "index": -1,
            }
        )
        access_token = "Bearer " + self.access_token
        headers = {
            "Authorization": access_token,
            "Content-Type": "application/json; charset=utf-8",
        }
        response = requests.request("POST", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        return response_json

    def update_blocks(self, block):
        block_id = block["block_id"]
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks/"
            + block_id
        )
        now = datetime.now()
        formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
        # print(f"formatted_time: {formatted_time}")
        block_element = block["text"]["elements"][0]
        block_element["text_run"]["content"] = (
            formatted_time + "【已通知】" + block_element["text_run"]["content"]
        )
        payload = json.dumps({"update_text_elements": {"elements": [block_element]}})

        access_token = "Bearer " + self.access_token
        headers = {
            "Authorization": access_token,
            "Content-Type": "application/json; charset=utf-8",
        }

        response = requests.request("PATCH", url, headers=headers, data=payload)
        if response.status_code == 200:
            response_json_data = response.json()
            response_json_element = response_json_data["data"]["block"]["text"][
                "elements"
            ][0]
            response_json_content = response_json_element["text_run"]["content"]
            print("updated block:", response_json_content)

    def check_notice_exists(self):
        now = datetime.now()
        current_date = now.date()
        # print(f"current_date: {current_date}")
        notice_filename = str(current_date) + " 通知.docx"
        notice_path = self.notice_dir / Path(notice_filename)
        self.notice_path = Path(notice_path)
        if not notice_path.exists():
            notice_template_year_month = current_date.strftime("%Y-%m")
            notice_template_path = (
                self.notice_dir / f"{notice_template_year_month}-xx 通知.docx"
            )
            if notice_template_path.exists():
                notice_path.write_bytes(notice_template_path.read_bytes())

    def deliver_and_reply_messages(self):
        def is_later_than_a_specific_time_yesterday(text_message_notified_time):
            now = datetime.now()
            yesterday = now - timedelta(days=1)
            time_object = datetime.strptime(self.str_specific_time_yesterday, "%H-%M")
            hours = time_object.hour
            minutes = time_object.minute
            specific_time_yesterday = datetime(
                yesterday.year, yesterday.month, yesterday.day, hours, minutes
            )
            yesterday_23_59 = datetime(
                yesterday.year, yesterday.month, yesterday.day, 23, 59, 59
            )
            return (
                specific_time_yesterday < text_message_notified_time <= yesterday_23_59
            )

        pending_message_blocks_list = []
        for message_block_index, message_block in enumerate(
            self.all_document_blocks[self.item_message_block_start_index :],
            start=self.item_message_block_start_index,
        ):
            if message_block["block_type"] == 2:
                message_block_element = message_block["text"]["elements"][0]
                if (
                    18 < len(message_block_element["text_run"]["content"])
                    and message_block_element["text_run"]["content"][19:24]
                    == "【已通知】"
                ):
                    text_message_notified_time = datetime.strptime(
                        message_block_element["text_run"]["content"][:19],
                        "%Y-%m-%d %H:%M:%S",
                    )
                    if is_later_than_a_specific_time_yesterday(
                        text_message_notified_time
                    ):
                        pending_message_blocks_list.append(message_block)
                        print("A message was left after 20:40 yesterday.")
                        continue
                    else:
                        pass
                elif message_block_element["text_run"]["content"] == "":
                    pass
                else:
                    pending_message_blocks_list.append(message_block)

        print("len(pending_message_blocks_list):", len(pending_message_blocks_list))
        if len(pending_message_blocks_list) != 0:
            self.check_notice_exists()
            notice = Document(self.notice_path)
            target_paragraph_text = self.notice_message_heading

            for paragraph_index, paragraph in enumerate(notice.paragraphs):
                if target_paragraph_text in paragraph.text:
                    for pending_message_block_index, pending_message_block in enumerate(
                        pending_message_blocks_list
                    ):
                        # Deliver messages
                        pending_message_block_element = pending_message_block["text"][
                            "elements"
                        ][0]
                        pending_message_element = pending_message_block_element[
                            "text_run"
                        ]["content"]
                        new_paragraph = notice.add_paragraph(pending_message_element)
                        # notice.paragraphs.insert(paragraph_index + 1, new_paragraph)
                        paragraph._p.addnext(new_paragraph._p)
                        notice.save(self.notice_path)
                        print(
                            f"delivered pending_message_element[{pending_message_block_index}]: {pending_message_element}"
                        )

                        # Reply messages
                        self.update_blocks(pending_message_block)

    def delete_document_children_blocks(self, document_children_block_index):
        # The feishu official development document is weird.
        parent_block_id = self.document_id
        url = (
            "https://open.feishu.cn/open-apis/docx/v1/documents/"
            + self.document_id
            + "/blocks/"
            + parent_block_id
            + "/children/batch_delete"
        )
        payload = json.dumps(
            {
                "start_index": document_children_block_index,
                "end_index": document_children_block_index + 1,
            }
        )
        access_token = "Bearer " + self.access_token
        headers = {
            "Authorization": access_token,
            "Content-Type": "application/json; charset=utf-8",
        }
        response = requests.request("DELETE", url, headers=headers, data=payload)
        # print("delete document children blocks response:", response.text)

    def delete_notified_messages(self):
        # Deletion is executed based on children blocks, not item blocks.

        def is_text_message_notified_24_hours_before(message_block_element):
            if (
                18 < len(message_block_element["text_run"]["content"])
                and message_block_element["text_run"]["content"][19:24] == "【已通知】"
            ):
                text_message_notified_time = datetime.strptime(
                    message_block_element["text_run"]["content"][:19],
                    "%Y-%m-%d %H:%M:%S",
                )
                # print(text_message_notified_time)
                current_time = datetime.now()
                time_difference = current_time - text_message_notified_time
                time_difference_seconds = time_difference.total_seconds()
                if 24 * 60 * 60 < time_difference_seconds:
                    return True
                else:
                    return False
            else:
                return False

        def obtain_blank_and_notified_message_blocks():
            deletion_dict = {}
            deletion_images_dict = {}

            # Process different dypes of blocks respectively for debugging.

            # Initialize a dict to record blank message block indexes.
            blank_message_blocks_dict = {}
            # Initialize a dict to record notified text messages block indexes.
            text_message_blocks_dict = {}
            notified_text_message_blocks_dict = {}
            # Initialize a dict to record all image indexes if all text messages are notified.
            image_message_blocks_dic = {}

            # There is no message_heading_text in the `self.message_blocks_list`.
            for children_message_block_index, message_block in enumerate(
                self.message_blocks_list, start=self.children_message_block_start_index
            ):
                if message_block["block_type"] == 2:
                    message_block_element = message_block["text"]["elements"][0]
                    if message_block_element["text_run"]["content"] == "":
                        # Record all blank message blocks
                        blank_message_blocks_dict.update(
                            {children_message_block_index: "blank message blocks"}
                        )
                    else:
                        # Record all text message blocks
                        text_message_blocks_dict.update(
                            {
                                children_message_block_index: message_block_element[
                                    "text_run"
                                ]["content"]
                            }
                        )
                        # Record all notified text message blocks
                        if is_text_message_notified_24_hours_before(
                            message_block_element
                        ):
                            notified_text_message_blocks_dict.update(
                                {
                                    children_message_block_index: message_block_element[
                                        "text_run"
                                    ]["content"]
                                }
                            )
                elif message_block["block_type"] == 27:
                    image_message_blocks_dic.update(
                        {children_message_block_index: "image message"}
                    )

            # Add blank message blocks to the deletion dict
            deletion_dict.update(blank_message_blocks_dict)
            # Add notified text message blocks to the deletion dict
            deletion_dict.update(notified_text_message_blocks_dict)
            # Add deletion image message blocks to the deletion dict
            if len(text_message_blocks_dict) == len(notified_text_message_blocks_dict):
                deletion_dict.update(deletion_images_dict)
            sorted_deletion_dict = dict(sorted(deletion_dict.items()))
            return sorted_deletion_dict

        if len(self.message_blocks_list) != 0:
            sorted_deletion_dict = obtain_blank_and_notified_message_blocks()
            if len(sorted_deletion_dict) != 0:
                # The dict does not fuction well.
                # As processed indexes may be the same.
                deletion_waiting_list = []
                for index, (key, value) in enumerate(sorted_deletion_dict.items()):
                    deletion_waiting_list.append({key - index: value})
                # Feishu server executes the deletion step by step
                for deletion_waiting_dict in deletion_waiting_list:
                    self.delete_document_children_blocks(
                        list(deletion_waiting_dict.keys())[0]
                    )
                    print("deletion_waiting_dict:", deletion_waiting_dict)


class FeishuBaseAPI:
    def __init__(
        self,
        notice_dir,
        notice_message_heading,
        app_id,
        app_secret,
        app_token,
        table_id,
        view_id,
        save_search_records_response_path_as_json=True,
    ) -> None:
        self.notice_dir = Path(notice_dir)
        self.notice_path = Path()
        script_dir = Path(__file__).resolve().parent
        self.notice_message_heading = notice_message_heading
        tenant_access_token = self.obtain_tenant_access_token(app_id, app_secret)
        self.access_token = tenant_access_token
        self.str_specific_time_yesterday = "16-00"

        """
        The Base information
        """
        self.app_token = app_token
        self.table_id = table_id
        self.view_id = view_id
        # The Base Records
        self.search_records_response_json = {}
        self.records = []
        self.undelivered_records = []
        # Records that need to be delivered the second time
        self.undelivered_records_2 = []

        self.preprocess_records()
        search_records_response_path = script_dir / "search_records_response.json"
        if save_search_records_response_path_as_json is True:
            with open(search_records_response_path, "w+", encoding="utf8") as f:
                json.dump(self.search_records_response_json, f, ensure_ascii=False)

    def obtain_tenant_access_token(self, app_id, app_secret):
        url = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal"
        headers = {
            "Content-Type": "application/json; charset=utf-8",
        }
        payload = json.dumps(
            {
                "app_id": app_id,
                "app_secret": app_secret,
            }
        )
        response = requests.request("POST", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        return response_json["tenant_access_token"]

    def search_records(self):
        url = (
            "https://open.feishu.cn/open-apis/bitable/v1/apps/"
            + self.app_token
            + "/tables/"
            + self.table_id
            + "/records/search"
        )
        access_token = "Bearer " + self.access_token
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": access_token,
        }
        payload = json.dumps(
            {
                "view_id": self.view_id,
            }
        )
        response = requests.request("POST", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        return response_json

    def update_a_record(self, record, payload):
        url = (
            "https://open.feishu.cn/open-apis/bitable/v1/apps/"
            + self.app_token
            + "/tables/"
            + self.table_id
            + "/records/"
            + record["record_id"]
        )
        access_token = "Bearer " + self.access_token
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": access_token,
        }
        response = requests.request("PUT", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        return response_json

    def check_notice_exists(self):
        now = datetime.now()
        current_date = now.date()
        # print(f"current_date: {current_date}")
        notice_filename = str(current_date) + " 通知.docx"
        notice_path = self.notice_dir / Path(notice_filename)
        self.notice_path = Path(notice_path)
        if not notice_path.exists():
            notice_template_year_month = current_date.strftime("%Y-%m")
            notice_template_path = (
                self.notice_dir / f"{notice_template_year_month}-xx 通知.docx"
            )
            if notice_template_path.exists():
                notice_path.write_bytes(notice_template_path.read_bytes())

    def preprocess_records(self):
        def is_later_than_a_specific_time_yesterday(first_notification_time):
            now = datetime.now()
            yesterday = now - timedelta(days=1)
            time_object = datetime.strptime(self.str_specific_time_yesterday, "%H-%M")
            hours = time_object.hour
            minutes = time_object.minute
            specific_time_yesterday = datetime(
                yesterday.year, yesterday.month, yesterday.day, hours, minutes
            )
            yesterday_23_59 = datetime(
                yesterday.year, yesterday.month, yesterday.day, 23, 59, 59
            )
            return specific_time_yesterday < first_notification_time <= yesterday_23_59

        self.search_records_response_json = self.search_records()
        self.records = self.search_records_response_json["data"]["items"]
        for record in self.records:
            if (
                "是否已通知" not in record["fields"]
                or record["fields"]["是否已通知"] is False
            ):
                self.undelivered_records.append(record)
            else:
                first_notification_timestamp_s = record["fields"]["通知时间"] / 1000
                first_notification_time = datetime.fromtimestamp(
                    first_notification_timestamp_s
                )
                if is_later_than_a_specific_time_yesterday(first_notification_time):
                    if (
                        "是否已第二次通知" not in record["fields"]
                        or record["fields"]["是否已第二次通知"] is False
                    ):
                        self.undelivered_records_2.append(record)

        print(f"{len(self.undelivered_records)} undelivered records")
        print(
            f"{len(self.undelivered_records_2)} record(s) need(s) to be undelivered for the second time"
        )

    def deliver_and_reply_messages(self):
        def deliver_and_reply_records(records, annotation):
            self.check_notice_exists()
            notice = Document(self.notice_path)
            target_paragraph_text = self.notice_message_heading
            for paragraph_index, paragraph in enumerate(notice.paragraphs):
                if target_paragraph_text in paragraph.text:
                    for record_index, record in enumerate(records):
                        # Skip the demo
                        if record["record_id"] == "recQRfdWut":
                            continue
                        # Deliver the message
                        timestamp_s = datetime.now().timestamp()
                        timestamp_ms = int(timestamp_s * 1000)
                        # message_text may have several lines
                        message_text = record["fields"]["留言内容"]
                        for message_text_line_index, message_text_line in enumerate(
                            message_text
                        ):
                            if annotation == "" and message_text_line["text"] != "\n":
                                new_paragraph = notice.add_paragraph(
                                    message_text_line["text"].replace("\n", "")
                                )
                                # notice.paragraphs.insert(paragraph_index + 1, new_paragraph)
                                paragraph._p.addnext(new_paragraph._p)
                            elif (
                                annotation == "The second time:"
                                and message_text_line["text"] != "\n"
                            ):
                                first_notification_timestamp_s = (
                                    record["fields"]["通知时间"] / 1000
                                )
                                first_notification_time = datetime.fromtimestamp(
                                    first_notification_timestamp_s
                                )
                                message_text_line_2 = (
                                    first_notification_time.strftime(
                                        "%Y-%m-%d %H:%M:%S"
                                    )
                                    + "【已通知】"
                                    + message_text_line["text"].replace("\n", "")
                                )
                                new_paragraph = notice.add_paragraph(
                                    message_text_line_2
                                )
                                # notice.paragraphs.insert(paragraph_index + 1, new_paragraph)
                                paragraph._p.addnext(new_paragraph._p)
                        notice.save(self.notice_path)
                        print(
                            f"{annotation} delivered record text[{record_index}]: {message_text}"
                        )

                        # Confirm the reply
                        if annotation == "":
                            payload = json.dumps(
                                {"fields": {"是否已通知": True, "通知时间": timestamp_ms}}  # type: ignore
                            )
                            self.update_a_record(record, payload)
                        else:
                            payload = json.dumps(
                                {"fields": {"是否已第二次通知": True, "第二次通知时间": timestamp_ms}}  # type: ignore
                            )
                            self.update_a_record(record, payload)

        if len(self.undelivered_records) != 0:
            deliver_and_reply_records(self.undelivered_records, "")
        if len(self.undelivered_records_2) != 0:
            deliver_and_reply_records(self.undelivered_records_2, "The second time:")

    def delete_a_record(self, record):
        url = (
            "https://open.feishu.cn/open-apis/bitable/v1/apps/"
            + self.app_token
            + "/tables/"
            + self.table_id
            + "/records/"
            + record["record_id"]
        )
        access_token = "Bearer " + self.access_token
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": access_token,
        }
        payload = ""
        response = requests.request("DELETE", url, headers=headers, data=payload)
        response_json = json.loads(response.text)
        if response_json["code"] == 0:
            print(f"Deleted record: {record["fields"]["留言内容"][0]["text"]}")
        return response_json

    def delete_notified_messages(self):
        def is_message_notified_36_hours_before(record):
            first_notification_timestamp_s = record["fields"]["通知时间"] / 1000
            first_notification_time = datetime.fromtimestamp(
                first_notification_timestamp_s
            )
            current_time = datetime.now()
            time_difference = current_time - first_notification_time
            time_difference_seconds = time_difference.total_seconds()
            if 36 * 60 * 60 < time_difference_seconds:
                return True
            else:
                return False

        for record in self.records:
            # Skip the demo
            if record["record_id"] == "recQRfdWut":
                continue
            if (
                "是否已通知" in record["fields"]
                and record["fields"]["是否已通知"] is True
            ):
                if is_message_notified_36_hours_before(record):
                    self.delete_a_record(record)


if __name__ == "__main__":
    debug_dev_document = False
    now = datetime.now()
    formatted_time = now.strftime("%Y-%m-%d %H:%M:%S")
    print(f"now: {formatted_time}")

    script_dir = Path(__file__).resolve().parent
    configuration_path = script_dir / "configuration.txt"
    config_patterns = {
        "notice_dir": r"notice_dir=\s*(.*)",
        "notice_message_heading": r"notice_message_heading=\s*(.*)",
        "resource": r"resource=\s*(.*)",
        "app_id": r"app_id=\s*(.*)",
        "app_secret": r"app_secret=\s*(.*)",
        "document_id": r"document_id=\s*(.*)",
        "dev_document_id": r"dev_document_id=\s*(.*)",
        "message_heading_text": r"message_heading_text=\s*(.*)",
        "app_token": r"app_token=\s*(.*)",
        "table_id": r"table_id=\s*(.*)",
        "view_id": r"view_id=\s*(.*)",
    }

    config = {key: None for key in config_patterns}

    with open(configuration_path, "r", encoding="utf-8") as file:
        for line in file:
            for key, pattern in config_patterns.items():
                match = re.search(pattern, line)
                if match and match.group(1):
                    config[key] = match.group(1).strip()

    if debug_dev_document and config["dev_document_id"]:
        config["document_id"] = config["dev_document_id"]
    elif not debug_dev_document and config["document_id"]:
        pass

    if config["resource"] is None:
        config["resource"] = "document"
    if config["app_id"] is None:
        config["app_id"] = "cli_a7cacfd2f43e100e"
    if config["app_secret"] is None:
        config["app_secret"] = "2Io4fqrt1fGeVu5Sh32MwxooMa8xM2pX"

    if config["resource"] == "document" or config["resource"] == "both":
        print("\nUsing Feishu Document:")
        feishu_docs_api = FeishuDocumentAPI(
            notice_dir=config["notice_dir"],
            notice_message_heading=config["notice_message_heading"],
            document_id=config["document_id"],
            app_id=config["app_id"],
            app_secret=config["app_secret"],
            message_heading_text=config["message_heading_text"],
            debug_offline_all_document_comments_response_json=False,
            save_all_document_blocks_response_as_json=True,
            debug_offline_all_document_blocks_response_json=False,
            save_all_document_comments_response_as_json=True,
        )
        feishu_docs_api.deliver_and_reply_messages()
        feishu_docs_api.delete_notified_messages()
    if config["resource"] == "base" or config["resource"] == "both":
        print("\nUsing Feishu Base:")
        feishu_base_api = FeishuBaseAPI(
            notice_dir=config["notice_dir"],
            notice_message_heading=config["notice_message_heading"],
            app_id=config["app_id"],
            app_secret=config["app_secret"],
            app_token=config["app_token"],
            table_id=config["table_id"],
            view_id=config["view_id"],
            save_search_records_response_path_as_json=True,
        )
        feishu_base_api.deliver_and_reply_messages()
        feishu_base_api.delete_notified_messages()
